"""Journal Formatting: compliance checklist + side-by-side rewrite + DOCX export.
Two journals in demo (per the risk decisions): Nature and IEEE."""
import io
import re

from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import StreamingResponse

from .. import tasks
from ..config import settings
from ..db import SessionLocal, get_db
from ..models import EvidenceItem, Section
from ..services import claude
from .manuscripts import get_ms

router = APIRouter()

# XML (and therefore DOCX) rejects control characters that survive PDF extraction.
_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean(text: str) -> str:
    return _CTRL.sub("", text or "")

JOURNALS = {
    "scientific-reports": {
        "name": "Scientific Reports",
        "article_type": "Research Article",
        "citation_style": "numbered, Nature style (superscript)",
        "abstract_max_words": 200,
        "section_order": ["Abstract", "Introduction", "Results", "Discussion",
                          "Methods", "Data Availability", "References"],
        "rules": [
            "Abstract <= 200 words, unreferenced, no subheadings",
            "Main text (excluding Methods/refs/captions) <= 4,500 words",
            "Methods placed after Discussion",
            "Data Availability Statement required",
            "Numbered citation style (superscript), max ~60 references",
            "Figures max 8; each with a title <= 15 words",
        ],
    },
    "ieee-access": {
        "name": "IEEE Access",
        "article_type": "Regular Paper",
        "citation_style": "numbered [1] IEEE style",
        "abstract_max_words": 250,
        "section_order": ["Abstract", "Introduction", "Related Work", "Methodology",
                          "Experiments", "Results", "Conclusion", "References"],
        "rules": [
            "Abstract <= 250 words, single paragraph, no references",
            "Index Terms (keywords, alphabetical) required after abstract",
            "Citations in square brackets [1], numbered by first appearance",
            "Section headings numbered with Roman numerals",
            "A Conclusion section is required",
            "All acronyms defined at first use",
        ],
    },
}

@router.get("/journals")
def list_journals():
    return [{"id": k, "name": v["name"], "article_type": v["article_type"],
             "rules": v["rules"]} for k, v in JOURNALS.items()]


def run_format(task_id: str, ms_id: str, journal_id: str):
    db = SessionLocal()
    try:
        journal = JOURNALS[journal_id]
        sections = (db.query(Section).filter_by(manuscript_id=ms_id)
                    .order_by(Section.order).all())
        content = "\n\n".join(
            f"## {s.name}\n{s.text[:3000]}" for s in sections
            if s.name.lower() not in ("references", "bibliography")
        )[:45000]

        tasks.update(task_id, step=f"Checking compliance vs {journal['name']} guidelines", progress=20)
        checklist = claude.complete_json(
            f"Check this manuscript against {journal['name']} submission rules:\n"
            + "\n".join(f"- {r}" for r in journal["rules"])
            + f"\nExpected section order: {journal['section_order']}\n\n"
            "Return JSON: {\"checks\": [{\"rule\": str, \"status\": \"pass\"|\"fail\"|\"warning\", "
            "\"detail\": \"one sentence, cite numbers (word counts etc.) when possible\"}]}\n\n"
            f"MANUSCRIPT:\n{content}",
            max_tokens=2500,
        )

        tasks.update(task_id, step="Reformatting abstract & structure", progress=55)
        rewrite = claude.complete_json(
            f"Reformat this manuscript for submission to {journal['name']} "
            f"(citation style: {journal['citation_style']}, "
            f"abstract max {journal['abstract_max_words']} words, "
            f"section order: {journal['section_order']}).\n"
            "Return JSON: {\"abstract_before\": \"current abstract (verbatim, may truncate to 150 words)\", "
            "\"abstract_after\": \"rewritten compliant abstract\", "
            "\"restructure_plan\": [{\"from\": \"current section\", \"to\": \"target section\", \"note\": str}], "
            "\"added_statements\": [\"e.g. Data Availability Statement draft\"]}\n\n"
            f"MANUSCRIPT:\n{content[:30000]}",
            model=settings.claude_model_smart, max_tokens=3000,
        )

        result = {"journal": journal["name"], "journal_id": journal_id,
                  "checklist": checklist.get("checks", []), "rewrite": rewrite}
        db.add(EvidenceItem(
            manuscript_id=ms_id,
            claim=f"Formatted for {journal['name']}: "
                  f"{sum(1 for c in result['checklist'] if c['status'] == 'pass')}"
                  f"/{len(result['checklist'])} rules pass",
            kind="format", source_type="manuscript_span",
            source_ref={"page": 1, "quote": "", "section": "Abstract"},
            confidence=0.9, status="verified",
        ))
        db.commit()
        tasks.finish(task_id, result)
    except Exception as e:
        tasks.fail(task_id, str(e)[:400])
    finally:
        db.close()


@router.post("/format/{ms_id}")
def start_format(ms_id: str, journal: str, background: BackgroundTasks, db=Depends(get_db)):
    get_ms(db, ms_id)
    if journal not in JOURNALS:
        raise HTTPException(400, f"journal must be one of {list(JOURNALS)}")
    task_id = tasks.create("format")
    background.add_task(run_format, task_id, ms_id, journal)
    return {"task_id": task_id}


@router.get("/format/{ms_id}/export")
def export_docx(ms_id: str, journal: str = "scientific-reports", db=Depends(get_db)):
    ms = get_ms(db, ms_id)
    j = JOURNALS.get(journal, JOURNALS["scientific-reports"])
    sections = {s.name: s.text for s in
                db.query(Section).filter_by(manuscript_id=ms_id).all()}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)
    doc.add_heading(_clean(ms.title), level=0)
    doc.add_paragraph(", ".join(ms.authors or []))
    note = doc.add_paragraph()
    note.add_run(f"Formatted for {j['name']} - generated by PaperClue").italic = True
    for name in j["section_order"]:
        text = None
        for sname, stext in sections.items():
            if sname.lower().startswith(name.lower()[:6]):
                text = stext
                break
        doc.add_heading(name, level=1)
        doc.add_paragraph(_clean(text[:8000]) if text else f"[{name} - to be completed]")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"{ms.title[:40].replace(' ', '_')}_{journal}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )
